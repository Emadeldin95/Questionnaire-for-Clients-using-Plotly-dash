import dash
from dash import dcc, html, Input, Output, State
import dash_bootstrap_components as dbc
import plotly.graph_objs as go
import pandas as pd
from docx import Document
import io
import base64

# Initialize the Dash app with Bootstrap styling
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.PULSE])

company_website = "https://www.alexdynamics.com"

# Header component of the application
header = dbc.Navbar(
    dbc.Container([
        # Logo linking to the company website
        html.A(html.Img(src="/assets/logo.png", id='header-logo', height="80px"), href=company_website, target="_blank"),
        # Title of the questionnaire
        html.H2("Robotic System Investment Questionnaire", className="pt-5 text-center mx-auto")
    ]),
    color="light", light=True,
    dark=True,
    className="mb-4"
)

# Footer component of the application
footer = dbc.Container([
    html.Hr(),
    dbc.Row([
        dbc.Col(html.P("© 2025 Alex Dynamics. All Rights Reserved.")),
        dbc.Col(html.P("info@alexdynamics.com")),
        dbc.Col(html.A(html.Img(src="/assets/logo.png", id='footer-logo', height="60px"), href=company_website, target="_blank"))
    ], className="text-center mt-3"),
    dbc.Row([
        dbc.Col(html.P(company_website, className="text-center"))
    ], className="text-center")
], fluid=True)

# Layout of the questionnaire
app.layout = dbc.Container([
    header,
    
    html.Div([
        html.Hr(),
        html.H3("Process Evaluation"),
        dbc.Row([
            # Task Type selection
            dbc.Col(dbc.Card([
                dbc.CardHeader("Task Type"),
                dbc.CardBody(dcc.RadioItems(id='task_type', options=[
                    {'label': 'Repetitive (+30)', 'value': 'Repetitive'},
                    {'label': 'Variable (-20)', 'value': 'Variable'},
                    {'label': 'Mixed (+10)', 'value': 'Mixed'}
                ], labelStyle={'display': 'block'}))
            ], className="mb-3"), width=6),
            
            # Annual Volume selection
            dbc.Col(dbc.Card([
                dbc.CardHeader("Annual Volume"),
                dbc.CardBody(dcc.RadioItems(id='annual_volume', options=[
                    {'label': 'Low (<50k units, -10)', 'value': 'Low'},
                    {'label': 'Medium (50k–200k, +15)', 'value': 'Medium'},
                    {'label': 'High (>200k, +30)', 'value': 'High'}
                ], labelStyle={'display': 'block'}))
            ], className="mb-3"), width=6)
        ]),
        
        # Pain Points selection
        dbc.Card([
            dbc.CardHeader("Pain Points"),
            dbc.CardBody(dcc.Checklist(id='pain_points', options=[
                {'label': 'Labor shortages (+10)', 'value': 'Labor'},
                {'label': 'Safety risks (+10)', 'value': 'Safety'},
                {'label': 'Quality issues (+10)', 'value': 'Quality'},
                {'label': 'Slow throughput (+10)', 'value': 'Throughput'}
            ], labelStyle={'display': 'inline-block', 'margin-right': '10px'}))
        ], className="mb-3"),
        
        # Technical Feasibility section
        html.H3("Technical Feasibility"),
        dbc.Card([
            dbc.CardHeader("Workspace Type"),
            dbc.CardBody(dcc.RadioItems(id='workspace_type', options=[
                {'label': 'Structured (+40)', 'value': 'Structured'},
                {'label': 'Semi-structured (+20)', 'value': 'Semi-structured'},
                {'label': 'Unstructured (-30)', 'value': 'Unstructured'}
            ], labelStyle={'display': 'inline-block', 'margin-right': '10px'}))
        ], className="mb-3"),
        
        # Financial Analysis section
        html.H3("Financial Analysis"),
        dbc.Card([
            dbc.CardHeader("Current Labor Cost"),
            dbc.CardBody([
                html.P("Enter the total annual labor cost associated with the current manual process."),
                dcc.Input(id='current_labor_cost', type='number', placeholder="Annual Labor Cost ($)", className="form-control mb-2")
            ])
        ], className="mb-3"),
        
        # Upfront Robot Cost input
        dbc.Card([
            dbc.CardHeader("Upfront Robot Cost"),
            dbc.CardBody([
                html.P("Enter the initial cost for purchasing and installing the robotic system."),
                dcc.Input(id='upfront_cost', type='number', placeholder="Upfront Robot Cost ($)", className="form-control mb-2")
            ])
        ], className="mb-3"),
        
        # Expected Throughput Gain input
        dbc.Card([
            dbc.CardHeader("Expected Throughput Gain"),
            dbc.CardBody([
                html.P("Estimate the percentage increase in production output due to automation."),
                dcc.Slider(id='throughput_gain', min=0, max=50, step=1, value=10,
                           marks={0: '0%', 50: '50%'}, tooltip={"placement": "bottom", "always_visible": True})
            ])
        ], className="mb-3"),
        
        # Submit Button
        html.Br(),
        dbc.Button("Calculate Results", id='calculate', color='danger', className="mb-3 btn-lg", style={"width": "50%", "display": "block", "margin": "auto", "borderRadius": "20px"}),
        
        # Output Section for results
        html.Div(id='output_results', className="alert alert-info"),
        html.Div(
            dcc.Graph(id='roi_chart', style={"width": "70%", "margin": "auto", "border": "2px solid #28a745", "padding": "15px", "borderRadius": "15px", "boxShadow": "5px 5px 15px rgba(0,0,0,0.2)"})
        ),
        
        # Download Report Button
        html.Br(),
        dbc.Button("Download Report", id='download_report', color='info', className="btn-lg", style={"width": "50%", "display": "block", "margin": "auto", "borderRadius": "20px"}),
        dcc.Download(id="download_file")
    ], style={"marginLeft": "4%", "marginRight": "4%"}),
    
    footer
], fluid=True)

# Callback to compute scores and generate the ROI chart
@app.callback(
    [Output('output_results', 'children'),
     Output('roi_chart', 'figure')],
    [Input('calculate', 'n_clicks')],
    [State('task_type', 'value'),
     State('annual_volume', 'value'),
     State('pain_points', 'value'),
     State('workspace_type', 'value'),
     State('current_labor_cost', 'value'),
     State('upfront_cost', 'value'),
     State('throughput_gain', 'value')]
)
def compute_scores(n_clicks, task_type, annual_volume, pain_points, workspace_type, labor_cost, upfront_cost, throughput_gain):
    if not n_clicks:
        return "", go.Figure()
    
    # Score Calculation
    process_score = 30 if task_type == 'Repetitive' else -20 if task_type == 'Variable' else 10
    process_score += 30 if annual_volume == 'High' else 15 if annual_volume == 'Medium' else -10
    process_score += len(pain_points) * 10 if pain_points else 0
    process_score = min(100, max(0, process_score))
    
    tech_score = 40 if workspace_type == 'Structured' else 20 if workspace_type == 'Semi-structured' else -30
    tech_score = min(100, max(0, tech_score))
    
    roi_5yr = ((labor_cost * 0.7) + (throughput_gain / 100 * labor_cost) - upfront_cost) / upfront_cost * 100 if upfront_cost else 0
    roi_5yr = min(100, max(0, roi_5yr))
    
    # Graph
    fig = go.Figure()
    fig.add_trace(go.Bar(name='Scores', x=['Process', 'Technical', 'ROI (5yr)'], y=[process_score, tech_score, roi_5yr]))
    fig.update_layout(title='Investment Decision Factors',
                      xaxis_title='Categories',
                      yaxis_title='Scores (%)',
                      barmode='group')
    
    result_text = f"Process Score: {process_score}%, Technical Score: {tech_score}%, ROI (5yr): {roi_5yr}%"
    return result_text, fig

# Callback to generate a report document
@app.callback(
    Output("download_file", "data"),
    [Input("download_report", "n_clicks")],
    [State('output_results', 'children')]
)
def generate_report(n_clicks, results_text):
    if not n_clicks:
        return None
    
    doc = Document()
    doc.add_heading("Robotic System Investment Report", level=1)
    doc.add_paragraph(results_text)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return dcc.send_bytes(file_stream.getvalue(), "investment_report.docx")

# Run the server
if __name__ == '__main__':
    app.run_server(debug=True)
